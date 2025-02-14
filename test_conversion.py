import logging
import os
from pathlib import Path
from file_converter import FileConverter

# Configure logging
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

def create_sample_docx():
    """Create a sample Word document for testing"""
    from docx import Document
    
    doc = Document()
    doc.add_paragraph('This is a test document.')
    doc.add_paragraph('It contains multiple paragraphs.')
    doc.add_paragraph('Used for testing file conversion.')
    
    file_path = 'temp/test_document.docx'
    doc.save(file_path)
    logger.info(f"Created sample DOCX at: {file_path}")
    return file_path

def create_sample_pdf():
    """Create a sample PDF for testing"""
    from fpdf import FPDF
    
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(0, 10, txt="This is a test PDF document.", ln=True)
    pdf.cell(0, 10, txt="Created for testing file conversion.", ln=True)
    
    file_path = 'temp/test_document.pdf'
    pdf.output(file_path)
    logger.info(f"Created sample PDF at: {file_path}")
    return file_path

def test_conversions():
    """Test both conversion directions"""
    converter = FileConverter()
    
    try:
        # Ensure temp directory exists
        temp_dir = Path('temp')
        temp_dir.mkdir(exist_ok=True)
        logger.info(f"Temp directory status - exists: {temp_dir.exists()}, is_dir: {temp_dir.is_dir()}")
        
        # Test Word to PDF conversion
        logger.info("Testing Word to PDF conversion...")
        docx_path = create_sample_docx()
        logger.info(f"Testing conversion of {docx_path}")
        if os.path.exists(docx_path):
            logger.info(f"DOCX file size: {os.path.getsize(docx_path)} bytes")
            pdf_path = converter.word_to_pdf(docx_path)
            logger.info(f"Successfully converted to PDF: {pdf_path}")
            logger.info(f"PDF file size: {os.path.getsize(pdf_path)} bytes")
        else:
            logger.error("Failed to create test DOCX file")
            
        # Test PDF to Word conversion
        logger.info("Testing PDF to Word conversion...")
        pdf_path = create_sample_pdf()
        logger.info(f"Testing conversion of {pdf_path}")
        if os.path.exists(pdf_path):
            logger.info(f"PDF file size: {os.path.getsize(pdf_path)} bytes")
            docx_path = converter.pdf_to_word(pdf_path)
            logger.info(f"Successfully converted to DOCX: {docx_path}")
            logger.info(f"DOCX file size: {os.path.getsize(docx_path)} bytes")
        else:
            logger.error("Failed to create test PDF file")
            
    except Exception as e:
        logger.error(f"Test failed with error: {str(e)}", exc_info=True)
    finally:
        # Cleanup test files
        for test_file in temp_dir.glob('test_document.*'):
            try:
                test_file.unlink()
                logger.info(f"Cleaned up test file: {test_file}")
            except Exception as e:
                logger.error(f"Failed to cleanup {test_file}: {str(e)}")

if __name__ == "__main__":
    test_conversions()
